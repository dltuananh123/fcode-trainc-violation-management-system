"""Update report.docx to reflect v2.3 changes."""
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
import re

DOC_PATH = r"C:\Users\Admin\Desktop\GIT CLONE edu\fcode-trainc-violation-management-system\report\report.docx"
doc = Document(DOC_PATH)

# ──────────────────────────────────────────────────────
# Helper: insert a new paragraph after a given paragraph
# ──────────────────────────────────────────────────────
def insert_paragraph_after(paragraph, text, style=None):
    """Insert a new paragraph after the given paragraph. Returns new paragraph."""
    new_p = docx.oxml.OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    run = new_para.add_run(text)
    return new_para

def set_run_text(cell, text):
    """Replace all text in a cell, preserving first run's formatting."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    if cell.paragraphs:
        p0 = cell.paragraphs[0]
        if p0.runs:
            p0.runs[0].text = text
        else:
            p0.add_run(text)
    else:
        p0 = cell.add_paragraph()
        p0.add_run(text)

# ──────────────────────────────────────────────────────
# 1. Update Table 2 (Project Period, Submission Date)
# ──────────────────────────────────────────────────────
t2 = doc.tables[2]
set_run_text(t2.rows[0].cells[1], "17 April 2026 - 1 June 2026")
set_run_text(t2.rows[3].cells[1], "1 June 2026")

# ──────────────────────────────────────────────────────
# 3. Section 1 text updates
# ──────────────────────────────────────────────────────
# P74: Password hashing (was "plain text" - now FNV-1a + XOR)
# Find the paragraph containing "Passwords are stored in plain text"
for p in doc.paragraphs:
    if "Passwords are stored in plain text" in p.text:
        p.runs[0].text = ""
        for r in p.runs[1:]:
            r.text = ""
        p.runs[0].text = (
            "Passwords are stored using salted FNV-1a hash with XOR obfuscation in accounts.dat. "
            "Each password is hashed with a unique per-account salt (10 rounds of FNV-1a), "
            "then the hash hex string is XOR-encrypted with the salt for storage. "
            "This prevents plain-text recovery: even with file access, credentials cannot be read directly."
        )
        break

# P75: ADMIN account issue - now fixed with SE203055
for p in doc.paragraphs:
    if "auto-created ADMIN account has no corresponding Member entry" in p.text:
        p.runs[0].text = ""
        for r in p.runs[1:]:
            r.text = ""
        p.runs[0].text = (
            "The default admin account (ADMIN/ADMIN) still has no Member entry, but the seed data "
            "now creates SE203055 (Nguyen Ngoc Phuc) as BCN account with full admin access and "
            "a corresponding Member record, so memberViewProfile() works correctly from BCN menu."
        )
        break

# P79: Update date
for p in doc.paragraphs:
    if "as of 16 May 2026" in p.text:
        for r in p.runs:
            if "16 May 2026" in r.text:
                r.text = r.text.replace("16 May 2026", "1 June 2026")
        break

# ──────────────────────────────────────────────────────
# 4. Update Table 4 (Feature Status) - update existing & add rows
# ──────────────────────────────────────────────────────
t4 = doc.tables[4]

# Row 5 (index 5): Add new member - update notes (auto-create account with default password = MSSV)
set_run_text(t4.rows[5].cells[4],
    "Default password = studentId; account auto-created; all "
    "edit fields re-prompt on validation failure; \"0 de quay lai\" supported"
)

# Row 6 (index 6): Edit member - update notes
set_run_text(t4.rows[6].cells[4],
    "Changing role recalculates unpaid fines; phone/email unique validation; "
    "name/phone/email fields wrapped in while(1) re-prompt loops"
)

# Row 8 (index 8): View own profile and member list - update notes
set_run_text(t4.rows[8].cells[4],
    "Member list supports pagination (15 rows/page, n/m/q navigation); "
    "5-column table: MSSV, Ho va ten, Email, SDT, Ban"
)

# Row 13 (index 13): View all violations with filter - update notes
set_run_text(t4.rows[13].cells[4],
    "Filters by team, reason, payment status; paginated with n/m/q; "
    "\"Cho dong phat\" column 15 chars, status column 12 chars"
)

# Row 17 (index 17): Search violations by date range - update notes
set_run_text(t4.rows[17].cells[4],
    "Uses normalized day boundaries; paginated with n/m/q; "
    "start/end date wrapped in outer while(1) retry loop"
)

# Row 18 (index 18): Persistent binary storage - update notes
set_run_text(t4.rows[18].cells[4],
    "Auto-creates ADMIN/ADMIN on first run; fileio.c uses .tmp/.bak strategy; "
    "passwords stored with FNV-1a hash + XOR encryption"
)

# Add new feature rows after row 20 (last row)
# New features to add:
# 21 - Kick/Restore member
# 22 - Discipline Dashboard
# 23 - Password hashing
# 24 - Email validation
# 25 - "0 de quay lai" (go back)
# 26 - Member list 5-column display

new_features = [
    ("21", "Kick (deactivate) and restore members",
     "Advanced", "Done",
     "Marks member as deleted with isDeleted=1, isActive=0; preserves violations; "
     "full restore reinstates member to active status"),

    ("22", "Discipline Dashboard (Out CLB tracking)",
     "Advanced", "Done",
     "Shows kicked members, members near Out CLB threshold, "
     "and recent violations in a consolidated dashboard view"),

    ("23", "Password hashing with FNV-1a + XOR encryption",
     "Advanced", "Done",
     "Salted FNV-1a hash with 10 rounds; per-account random salt; "
     "XOR encryption of hash hex string; prevents plain-text recovery"),

    ("24", "Input validation enhancements",
     "Required", "Done",
     "Email: rejects commas, invalid TLD chars, dots at start/end, "
     "consecutive dots; enforces min 2-letter TLD. Phone: unique validation. "
     "All edit fields re-prompt on failure"),

    ("25", "\"0 de quay lai\" (zero to go back) on MSSV prompts",
     "Required", "Done",
     "All student ID prompts support entering 0 to cancel and return "
     "to previous menu; implemented in member/violation operations"),

    ("26", "Pagination across list views",
     "Advanced", "Done",
     "Member list, violation view (all/filtered), and date search "
     "use consistent clear-screen paging: 15 rows/page, n/m/q navigation"),
]

for feat in new_features:
    row = t4.add_row()
    for i, text in enumerate(feat):
        set_run_text(row.cells[i], text)

# ──────────────────────────────────────────────────────
# 5. Update Table 5 (Directory Structure)
# ──────────────────────────────────────────────────────
t5 = doc.tables[5]
# Add rows for run.bat and report/
row = t5.add_row()
set_run_text(row.cells[0], "run.bat")
set_run_text(row.cells[1],
    "One-click build-and-run script. Calls mingw32-make, compiles seed_data.exe, "
    "runs seed, then launches the application."
)
# Add report/ row
row2 = t5.add_row()
set_run_text(row2.cells[0], "report/")
set_run_text(row2.cells[1],
    "Defense report template (report.docx) and update script."
)

# Update existing rows to reflect new counts
set_run_text(t5.rows[8].cells[1],
    "Standalone generator for deterministic seed data: 72 members (70 real Challenge 3 students "
    "+ SE203055 + ADMIN), 72 accounts, 76 violations, 3 kicked members, 1 OUT_CLB case."
)
set_run_text(t5.rows[9].cells[1],
    "Repository-level binary seed files: accounts.dat, members.dat, violations.dat with real data."
)

# ──────────────────────────────────────────────────────
# 6. Update Table 6 (Account struct) - password field
# ──────────────────────────────────────────────────────
t6 = doc.tables[6]
# Row 5 (password field) - change description
set_run_text(t6.rows[5].cells[2],
    "Salted FNV-1a hash (10 rounds) with per-account random salt, "
    "XOR-encrypted with the salt for storage. Not recoverable as plain text."
)

# ──────────────────────────────────────────────────────
# 7. Update Table 11 (Demo Walkthrough)
# ──────────────────────────────────────────────────────
t11 = doc.tables[11]
# Update demo steps to reflect new seed data
demo_steps = [
    ("1", "Build and seed data",
     "Run mingw32-make, then run bin\\seed_data.exe or use run.bat"),
    ("2", "Program startup",
     "Run bin\\violation-management-system.exe. Login screen appears."),
    ("3", "Admin login (BCN)",
     "Login with SE203055 / Phuc@2006 (BCN role). The BCN menu shows all "
     "management options including member CRUD, violation recording, kick/restore."),
    ("4", "View member list",
     "Choose option 13. Shows 5-column table (MSSV, Ho va ten, Email, SDT, Ban) "
     "with pagination. Use n=next page, m=previous page, q=quit."),
    ("5", "Add a new member",
     "Choose option 1, enter a test student ID (e.g., SE209999). All fields validated; "
     "invalid input re-prompts until corrected. Enter 0 at MSSV prompt to cancel."),
    ("6", "Record a violation",
     "Choose option 4, enter a member's MSSV (e.g., SE204111 for Anh). "
     "Select reason and confirm. Fine calculated automatically by role."),
    ("7", "Check Out CLB threshold",
     "Choose option 8. Shows members at risk; 4+ consecutive absences prompts "
     "BCN to confirm Out CLB action."),
    ("8", "Kick and restore member",
     "Choose option 15 to kick a member (e.g., SE210946, SE210117, SE203367). "
     "Choose option 16 to restore. Verify status change in member list."),
    ("9", "Filtered violation list",
     "Choose option 6, then filter by team/reason/payment. Paginated results."),
    ("10", "Mark fine as paid",
     "Choose option 5, enter a member MSSV. View unpaid violations, mark as paid."),
    ("11", "Statistics by team",
     "Choose option 7. Shows collected, outstanding, and total fines per team."),
    ("12", "Sort by violation count",
     "Choose option 9, then descending order. Members sorted by violation count."),
    ("13", "Export report",
     "Choose option 10. Generates a .txt report beside the executable."),
    ("14", "Search by date range",
     "Choose option 11. Enter valid date range; paginated results. "
     "Invalid dates re-prompt until corrected."),
    ("15", "Member self-service",
     "Logout and login as a member (any student MSSV as password, e.g., "
     "SE204111/SE204111). Member menu shows profile, own violations, own fines."),
    ("16", "Discipline Dashboard",
     "BCN option 14. Shows kicked members, near-threshold warnings, recent violations."),
    ("17", "Persistence check",
     "Exit and reopen. All data persists: members, violations, payment status, kicks."),
]
# Clear existing rows (keep header at index 0) and re-add
# Remove all rows after header
while len(t11.rows) > 1:
    row = t11.rows[-1]
    tbl_element = t11._tbl
    tbl_element.remove(row._tr)

for step in demo_steps:
    row = t11.add_row()
    for i, text in enumerate(step):
        set_run_text(row.cells[i], text)

# ──────────────────────────────────────────────────────
# 8. Update Table 12 (Known Issues)
# ──────────────────────────────────────────────────────
t12 = doc.tables[12]
# Remove the two known issues (password plain text, admin no member)
# and replace with updated issues
# First, remove all rows after header
while len(t12.rows) > 1:
    row = t12.rows[-1]
    tbl_element = t12._tbl
    tbl_element.remove(row._tr)

# Add updated issues
known_issues = [
    ("Invalid calendar dates may be accepted",
     "parseDate() uses mktime() after sscanf(), which silently normalizes "
     "invalid dates (e.g., 31/02/2026 becomes 03/03/2026). Users may not "
     "realize the date was adjusted.",
     "High",
     "Validate day/month/year combinations before calling mktime()."),
    ("Violation record lookup by name localized",
     "violationRecord() supports lookup by name, but diacritic characters "
     "in Vietnamese names may cause mismatches.",
     "Low",
     "Consider adding diacritic-normalization or fuzzy matching."),
    ("ADMIN legacy account has no Member record",
     "The auto-created ADMIN/ADMIN account from first-run bootstrap has no "
     "corresponding Member entry. Accessing profile view from this account fails.",
     "Low",
     "Use SE203055 (Phuc@2006) for BCN access; ADMIN account kept for backward "
     "compatibility only."),
    ("Fixed-width columns may truncate long names",
     "Member list uses fixed-width truncation (%-10.10s) for safety. "
     "Some Vietnamese names with diacritics may appear cut off.",
     "Low",
     "Consider dynamic column sizing or wider default widths."),
]
for issue in known_issues:
    row = t12.add_row()
    for i, text in enumerate(issue):
        set_run_text(row.cells[i], text)

# ──────────────────────────────────────────────────────
# 9. Update Table 14 (Repository URL) if needed
# ──────────────────────────────────────────────────────
t14 = doc.tables[14]
# Keep as-is, URL may be different but user didn't specify change

# ──────────────────────────────────────────────────────
# 10. Update paragraph P100 (directory structure observation)
# ──────────────────────────────────────────────────────
for p in doc.paragraphs:
    if "includes both engineering documents" in p.text:
        p.runs[0].text = ""
        for r in p.runs[1:]:
            r.text = ""
        p.runs[0].text = (
            "The repository includes both engineering documents (architecture.md, epics.md, story files), "
            "operational documents (demo-and-test-guide.md, technical_updates.md), "
            "and a defense report (report/report.docx)."
        )
        break

# ──────────────────────────────────────────────────────
# 11. Add new algorithm sections (3.3.10, 3.3.11)
# ──────────────────────────────────────────────────────
# Find the paragraph after P157 (end of 3.3.9) 
# We'll find it by content
target_p = None
for i, p in enumerate(doc.paragraphs):
    if "deterministic presentation path for defense day" in p.text:
        target_p = p
        break

if target_p:
    # Insert empty paragraph for spacing
    p_spacer1 = insert_paragraph_after(target_p, "", style=target_p.style)
    
    # 3.3.10 Password hashing
    p_title = insert_paragraph_after(p_spacer1, "3.3.10.  Password hashing with FNV-1a + XOR encryption", style="Heading 3")
    p_body = insert_paragraph_after(p_title, "", style="Normal")
    p_body.runs[0].text = (
        "Passwords are not stored as plain text. The auth module hashes each password using FNV-1a "
        "(Fowler-Noll-Vo) hash with a per-account random salt. The hash undergoes 10 rounds of FNV-1a "
        "with the salt XOR-mixed at each round. The final 64-bit hash is converted to a hex string, "
        "then XOR-encrypted with the salt bytes for storage. On login, the stored salt is used to "
        "recompute the hash from the entered password and compare against the stored obfuscated value. "
        "This ensures that even if accounts.dat is compromised, passwords cannot be recovered as plain text."
    )
    
    # 3.3.11 Member kick/restore
    p_title2 = insert_paragraph_after(p_body, "3.3.11.  Member kick (deactivation) and restore", style="Heading 3")
    p_body2 = insert_paragraph_after(p_title2, "", style="Normal")
    p_body2.runs[0].text = (
        "The system supports soft-deactivation of members through a kick operation. "
        "Kicking a member sets isDeleted=1, isActive=0, records a deletedAt timestamp, "
        "and preserves all violation history. The corresponding account is locked to prevent login. "
        "Restore reverses the operation: sets isDeleted=0, isActive=1, clears deletedAt, "
        "and unlocks the account. This allows the club to track former members without "
        "permanently removing their records. The discipline dashboard provides visibility "
        "into currently kicked members."
    )
    
    # 3.3.12 Input validation and re-prompt loops
    p_title3 = insert_paragraph_after(p_body2, "3.3.12.  Input validation and re-prompt loops", style="Heading 3")
    p_body3 = insert_paragraph_after(p_title3, "", style="Normal")
    p_body3.runs[0].text = (
        "All mutation operations use defensive while(1) loops that re-prompt on validation failure. "
        "If the user presses Enter without input, the existing value is kept. If input is invalid, "
        "an error message is shown and the prompt repeats. All student ID prompts support \"0 de quay lai\" "
        "(zero to go back) to cancel the operation and return to the previous menu. "
        "Email validation is particularly strict: it rejects commas, invalid TLD characters, "
        "dots at the start or end of the local part, consecutive dots, and enforces a minimum "
        "2-letter TLD. Phone numbers are normalized to 10 digits and checked for uniqueness."
    )
    
    # 3.3.13 Pagination
    p_title4 = insert_paragraph_after(p_body3, "3.3.13.  Pagination for list views", style="Heading 3")
    p_body4 = insert_paragraph_after(p_title4, "", style="Normal")
    p_body4.runs[0].text = (
        "Member list, violation view (all/filtered), and date search all use a consistent "
        "pagination system. Each view displays 15 rows per page (ROWS_PER_PAGE from include/ui.h). "
        "The user navigates with n (next page), m (previous page), and q (quit back to menu). "
        "The screen is cleared between pages for clean presentation. Violation list uses a "
        "two-pass approach: the first pass collects matching indices into a matchIdx[] array, "
        "and the second pass renders the current page from that index array."
    )

# ──────────────────────────────────────────────────────
# 12. Save
# ──────────────────────────────────────────────────────
doc.save(DOC_PATH)
print("report.docx updated successfully!")
